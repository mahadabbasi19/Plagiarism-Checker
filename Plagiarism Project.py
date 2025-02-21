from docx import Document
from PyPDF2 import PdfReader
from rapidfuzz import fuzz



def plagiarism_Detector(file1, file2):

    # CHECKING FOR TXT FILE
    if ".txt" in file1 and ".txt" in file2:

        # file 1
        x = open(file1, 'r')
        text1 = x.read()

        # file 2
        y = open(file2, 'r')
        text2 = y.read()

    # CHECKING FOR DOCX
    elif '.docx' in file1 and '.docx' in file2:
        # file 1
        document1 = Document(file1)
        text1 = ""
        for paragraph in document1.paragraphs:
            text1 += paragraph.text
        # file2
        document2 = Document(file2)
        text2 = ""
        for paragraph in document2.paragraphs:
            text2 += paragraph.text

    # CHECKING FOR PDF FILE
    elif '.pdf' in file1 and '.pdf' in file2:

        # pdf 1
        pdf1 =PdfReader(file1)
        text1 = ""
        for i in range(len(pdf1.pages)):
            page = pdf1.pages[i]
            text1 += page.extract_text()

        # pdf 2
        pdf2 = PdfReader(file1)
        text2 = ""
        for j in range(len(pdf2.pages)):
            page2 = pdf2.pages[j]
            text2 += page.extract_text()

    similarity_ratio =fuzz.ratio(text1, text2)
    if similarity_ratio >= 70:
        return f"Similarity Percent: {round(similarity_ratio)}%\nPLAGIARISM ALERT!!"
    return f"Similarity Percent: {round(similarity_ratio)}%\nNO PLAGIARISM FOUND."



file_1,file_2=input("Enter Names of Two Files(Comma Seprated): ").split(',')

print(plagiarism_Detector(file_1, file_2))
